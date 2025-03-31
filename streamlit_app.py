import streamlit as st
import requests
import cohere
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os
import markdown
from bs4 import BeautifulSoup

# --- CONFIGURATION ---
GEMINI_KEY = st.secrets["GEMINI_KEY"]
COHERE_KEY = st.secrets["COHERE_KEY"]
co = cohere.Client(COHERE_KEY)

# --- UTILITY FUNCTIONS ---
def markdown_to_text(md_string):
    html = markdown.markdown(md_string)
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text()

def simulate_financials(initial_users, fee, churn, growth_rate, fixed_cost, var_cost, months):
    users = [initial_users]
    revenue, total_cost, net_profit, cash_flow = [], [], [], []
    cumulative = 0
    for m in range(1, months+1):
        if m > 1:
            new_users = users[-1] * growth_rate
            churn_users = users[-1] * churn
            current_users = users[-1] + new_users - churn_users
            users.append(current_users)
        else:
            current_users = users[-1]
        monthly_revenue = current_users * fee
        monthly_cost = fixed_cost + (current_users * var_cost)
        profit = monthly_revenue - monthly_cost
        cumulative += profit
        revenue.append(monthly_revenue)
        total_cost.append(monthly_cost)
        net_profit.append(profit)
        cash_flow.append(cumulative)
    df = pd.DataFrame({
        "Month": list(range(1, months+1)),
        "Users": users,
        "Revenue": revenue,
        "Total Cost": total_cost,
        "Net Profit": net_profit,
        "Cash Flow": cash_flow
    })
    return df

def generate_financial_graph(df):
    plt.figure(figsize=(8, 5))
    plt.plot(df["Month"], df["Cash Flow"], marker="o", linestyle="-", color="blue")
    plt.title("Financial Projection")
    plt.xlabel("Month")
    plt.ylabel("Cumulative Cash Flow")
    plt.grid()
    st.pyplot(plt)

def generate_section(section_title, context):
    prompt = f"""You are a top-tier startup consultant.
Write a detailed section titled "{section_title}" based on the following context:
{context}
Ensure the analysis is insightful, data-driven, and consultative.
Write only the content for the section."""
    
    response = co.chat(
        model="command-xlarge-nightly",
        message=prompt,
        max_tokens=2189
    )
    return response.text.strip()

def build_startup_report(bplan, personality_type, psychometric, financial_params):
    doc = Document()
    doc.add_heading("Startup Consultant Report", level=1)
    
    sections = {
        "1. Overview of Input B-Plan": generate_section("Overview of Input B-Plan", bplan),
        "2. Founder-Market Fit": generate_section("Founder-Market Fit", f"{bplan}\n{personality_type}\n{psychometric}"),
        "3. Total Addressable Market (TAM)": generate_section("Total Addressable Market (TAM)", bplan),
        "4. Product-Market Fit": generate_section("Product-Market Fit", bplan),
        "5. Market Research Plan": generate_section("Market Research Plan", bplan)
    }
    
    df = simulate_financials(
        float(financial_params["initial_users"]),
        float(financial_params["fee"]),
        float(financial_params["churn"]),
        float(financial_params["growth_rate"]),
        float(financial_params["fixed_cost"]),
        float(financial_params["var_cost"]),
        int(financial_params["months"])
    )
    
    for title, content in sections.items():
        doc.add_heading(title, level=2)
        doc.add_paragraph(markdown_to_text(content))
    
    output_filename = "Startup_Consultant_Report.docx"
    doc.save(output_filename)
    return output_filename, df

# --- STREAMLIT APP UI ---
st.title("Startup Consultant Report Generator")

bplan = st.text_area("Enter your Business Plan")
personality_type = st.text_input("Founder Personality Type")
psychometric = st.text_area("Psychometric Insights")

st.subheader("Financial Parameters")
initial_users = st.number_input("Initial Users", min_value=1, value=100)
fee = st.number_input("Revenue per User", min_value=0.0, value=10.0)
churn = st.number_input("Churn Rate", min_value=0.0, max_value=1.0, value=0.1)
growth_rate = st.number_input("Growth Rate", min_value=0.0, max_value=1.0, value=0.2)
fixed_cost = st.number_input("Fixed Monthly Cost", min_value=0.0, value=1000.0)
var_cost = st.number_input("Variable Cost per User", min_value=0.0, value=2.0)
months = st.number_input("Months", min_value=1, value=12)

if st.button("Generate Report"):
    financial_params = {
        "initial_users": initial_users,
        "fee": fee,
        "churn": churn,
        "growth_rate": growth_rate,
        "fixed_cost": fixed_cost,
        "var_cost": var_cost,
        "months": months
    }
    report_filename, df = build_startup_report(bplan, personality_type, psychometric, financial_params)
    st.success(f"Report generated: {report_filename}")
    st.download_button(
        label="Download Report",
        data=open(report_filename, "rb").read(),
        file_name=report_filename
    )
    st.subheader("Financial Graph")
    generate_financial_graph(df)
